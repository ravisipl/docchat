from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangChainDocument  
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_community.vectorstores import PGVector
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from sqlalchemy.orm import Session
from sqlalchemy import text
from ..models.document import Document, DocumentChunk
from ..config import settings
import numpy as np
import os
from docx import Document as DocxDocument
from .embedding_base import EmbeddingProvider
from .embedding_gemini import GeminiEmbeddingProvider
from .embedding_openai import OpenAIEmbeddingProvider
from .embedding_huggingface import HuggingFaceEmbeddingProvider

def extract_text_from_pdf(file_path: str) -> list:
    """Extract text and metadata from PDF file"""
    loader = PyMuPDFLoader(file_path)
    documents = loader.load()
    return documents

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def get_text_content(file_path: str, file_name: str) -> list:
    """Extract text content based on file type"""
    file_ext = os.path.splitext(file_name)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text = extract_text_from_docx(file_path)
        return [LangChainDocument(page_content=text, metadata={"source": file_path, "filename": file_name})]
    elif file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return [LangChainDocument(page_content=text, metadata={"source": file_path, "filename": file_name})]
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def create_document_chunks(documents: list, file_path: str, file_name: str) -> list:
    """Create document chunks from text content"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    # For PDFs, we already have page-wise documents with metadata
    if any('page' in doc.metadata for doc in documents):
        return text_splitter.split_documents(documents)
    
    # For other file types, create chunks from the single document
    return text_splitter.split_documents(documents)

def get_embedding_provider():
    provider = os.getenv("EMBEDDING_PROVIDER", "gemini")
    if provider == "openai":
        return OpenAIEmbeddingProvider(api_key=settings.OPENAI_API_KEY)
    elif provider == "huggingface":
        return HuggingFaceEmbeddingProvider(model_name=settings.HF_API_KEY)
    else:
        return GeminiEmbeddingProvider(api_key=settings.GEMINI_API_KEY)

def initialize_embeddings() -> EmbeddingProvider:
    return get_embedding_provider()

def create_document_record(file_name: str, file_path: str, user_id: int, db: Session) -> Document:
    """Create and store document record in database"""
    db_document = Document(
        title=file_name,
        file_path=file_path,
        uploaded_by=user_id
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    return db_document

def store_chunks_with_embeddings(
    chunks: list,
    document_id: int,
    embeddings: EmbeddingProvider,
    db: Session
) -> None:
    """Store document chunks with their embeddings"""
    for i, chunk in enumerate(chunks):
        embedding_vector = embeddings.embed_query(chunk.page_content)
        embedding_list = embedding_vector.tolist() if isinstance(embedding_vector, np.ndarray) else list(embedding_vector)
        
        db_chunk = DocumentChunk(
            document_id=document_id,
            chunk_text=chunk.page_content,
            chunk_index=i,
            embedding=embedding_list
        )
        db.add(db_chunk)
    
    db.commit()

def store_in_pgvector(chunks: list, embeddings: EmbeddingProvider, collection_name: str) -> None:
    """Store documents in PGVector for similarity search"""
  
    PGVector.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=settings.VECTOR_STORE_COLLECTION,
        connection_string=settings.DATABASE_URL,
        pre_delete_collection=False
    )

async def process_document(file_path: str, file_name: str, user_id: int, db: Session):
    """
    Process a document file (TXT, PDF, or DOCX):
    1. Load and split the document into chunks
    2. Create embeddings using Gemini API
    3. Store document and chunks in PostgreSQL with vector embeddings
    """
    try:
        # Extract text content with metadata
        documents = get_text_content(file_path, file_name)
        
        # Create document chunks
        chunks = create_document_chunks(documents, file_path, file_name)
        
        # Initialize embeddings
        embeddings = initialize_embeddings()
        
        # Create document record
        db_document = create_document_record(file_name, file_path, user_id, db)
        
        # Ensure pgvector extension is enabled
        db.execute(text('CREATE EXTENSION IF NOT EXISTS vector'))
        
        # Store chunks with embeddings
        store_chunks_with_embeddings(chunks, db_document.id, embeddings, db)
        
        # Store in PGVector
        store_in_pgvector(chunks, embeddings, settings.VECTOR_STORE_COLLECTION)
        
        return db_document
        
    except Exception as e:
        db.rollback()
        raise e 